from __future__ import annotations

import json
import os
import re
import secrets
import time
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from core.base_docs_loader import BaseDocsLoader
from core.docx_parser import DocxParser
from core.security import sanitize_filename
from core.simple_validator import assess_lesson_text, normalize_lesson_text


REQUIRED_MAIN_SECTIONS = (
    "Тема уроку",
    "Клас",
    "Предмет",
    "Тип уроку",
    "Мета уроку",
    "Очікувані результати",
    "Обладнання та матеріали",
    "Хід уроку",
    "Домашнє завдання",
)

REQUIRED_FLOW_STAGES = (
    "Організаційний момент",
    "Актуалізація опорних знань",
    "Мотивація навчальної діяльності",
    "Вивчення нового матеріалу",
    "Закріплення знань",
    "Підсумок уроку",
)


class LessonGeneratorV25:
    """Unified full-lesson generator.

    v2.5 deliberately avoids assembling a lesson from tiny model calls. The
    model receives a parsed template blueprint and returns one complete lesson
    JSON document. Local code validates structure and renders the final DOCX.
    """

    pipeline_version = "generator_v2.5"

    def __init__(self, legacy_generator: Any):
        self.legacy = legacy_generator
        self.loader: BaseDocsLoader = getattr(legacy_generator, "loader", None) or BaseDocsLoader()

    @staticmethod
    def _clean(value: Any) -> str:
        return re.sub(r"\s+", " ", str(value or "")).strip()

    @staticmethod
    def _topic_tokens(value: str) -> set[str]:
        tokens = re.findall(r"[A-Za-z\u0400-\u04FF0-9']+", str(value or "").lower())
        stop = {"урок", "тема", "клас", "предмет", "числа", "число"}
        return {token for token in tokens if len(token) >= 3 and token not in stop}

    @staticmethod
    def _extract_target_number(value: str) -> str:
        text = str(value or "").lower()
        for pattern in (
            r"(?:множ|добут|ділен|частк)\w*[^.\n]{0,24}?\bна\s+(\d{1,2})\b",
            r"\bна\s+(\d{1,2})\b[^.\n]{0,24}?(?:множ|ділен)\w*",
        ):
            match = re.search(pattern, text)
            if match:
                return match.group(1)
        return ""

    @staticmethod
    def _stage_key(value: str) -> str:
        text = str(value or "").lower()
        if any(token in text for token in ("організа", "готовн")):
            return "organize"
        if any(token in text for token in ("актуал", "опорн", "повтор")):
            return "actualize"
        if any(token in text for token in ("мотива", "зацікав", "проблем")):
            return "motivate"
        if any(token in text for token in ("нов", "вивчен", "пояснен")):
            return "new_material"
        if any(token in text for token in ("закріп", "тренув", "практич")):
            return "practice"
        if any(token in text for token in ("підсум", "рефлекс", "оцін")):
            return "summary"
        return re.sub(r"\s+", " ", text).strip()

    @classmethod
    def _topic_kind(cls, value: str) -> str:
        text = str(value or "").lower()
        has_mult = "множ" in text or "добут" in text
        has_div = "ділен" in text or "частк" in text
        has_geometry = any(token in text for token in ("периметр", "площа", "прямокут", "геометр"))
        has_language = any(token in text for token in ("реченн", "слово", "текст", "іменник", "прикметник", "дієслов"))
        has_yads = any(token in text for token in ("дослід", "спостереж", "природ", "вода", "рослин", "тварин"))
        if has_geometry:
            return "geometry"
        if has_mult and has_div:
            return "math_mult_div"
        if has_mult:
            return "math_mult"
        if has_div:
            return "math_div"
        if has_language:
            return "language"
        if has_yads:
            return "yads"
        return "general"

    @staticmethod
    def _extract_reference_topic(parsed_doc: dict, path: Path) -> str:
        for field in parsed_doc.get("header_fields") or []:
            label = str(field.get("label") or "").strip().lower()
            if label in {"тема", "тема уроку"}:
                value = str(field.get("value") or "").strip()
                if value:
                    return value
        return path.stem

    @classmethod
    def _score_template(cls, *, topic: str, parsed_doc: dict, path: Path) -> float:
        reference_topic = cls._extract_reference_topic(parsed_doc, path)
        topic_tokens = cls._topic_tokens(topic)
        ref_tokens = cls._topic_tokens(reference_topic)
        overlap = len(topic_tokens & ref_tokens)
        topic_kind = cls._topic_kind(topic)
        ref_kind = cls._topic_kind(reference_topic)
        score = overlap * 12.0
        score += min(8, len(parsed_doc.get("stages") or [])) * 1.5
        score += min(6, len(parsed_doc.get("sections") or [])) * 0.8
        if topic_kind != "general" and topic_kind == ref_kind:
            score += 18.0
        elif topic_kind != "general" and ref_kind != "general":
            score -= 12.0

        target_number = cls._extract_target_number(topic)
        ref_numbers = set(re.findall(r"\b\d{1,2}\b", reference_topic))
        if target_number and ref_numbers and target_number not in ref_numbers:
            score -= 35.0
        return score

    def _select_template(self, *, subject: str, grade: str, topic: str) -> tuple[Path | None, dict]:
        paths = self.loader.load_docs_for_subject(subject, grade, limit=12)
        scored: list[tuple[float, Path, dict]] = []
        for path in paths:
            try:
                parsed = DocxParser.extract_structure(path)
            except Exception:
                continue
            scored.append((self._score_template(topic=topic, parsed_doc=parsed, path=path), path, parsed))
        if not scored:
            return None, {}
        scored.sort(key=lambda item: item[0], reverse=True)
        _, path, parsed = scored[0]
        return path, parsed

    @staticmethod
    def _sample_items(items: list[Any], limit: int = 6, max_chars: int = 900) -> list[str]:
        result: list[str] = []
        total = 0
        for item in items or []:
            text = re.sub(r"\s+", " ", str(item or "")).strip()
            if not text:
                continue
            if len(text) > 220:
                text = text[:220].rstrip() + "..."
            if total + len(text) > max_chars:
                break
            result.append(text)
            total += len(text)
            if len(result) >= limit:
                break
        return result

    def _build_blueprint(self, parsed_doc: dict, template_path: Path | None, *, subject: str) -> dict:
        header_labels = [
            self._clean(field.get("label"))
            for field in parsed_doc.get("header_fields") or []
            if self._clean(field.get("label"))
        ]
        sections = []
        for section in parsed_doc.get("sections") or []:
            title = self._clean(section.get("title") or section.get("display_title"))
            if title:
                sections.append(
                    {
                        "title": title,
                        "sample_items": self._sample_items(section.get("items") or [], limit=4),
                    }
                )
        stages = []
        for stage in parsed_doc.get("stages") or []:
            title = self._clean(stage.get("title") or stage.get("display_title"))
            if title:
                stages.append(
                    {
                        "title": title,
                        "sample_items": self._sample_items(stage.get("items") or [], limit=5),
                    }
                )
        return {
            "template_file": template_path.name if template_path else "",
            "subject_family": self.loader.resolve_subject(subject).get("canonical") or "",
            "header_labels": header_labels[:12],
            "sections": sections[:10],
            "stages": stages[:12],
            "required_main_sections": list(REQUIRED_MAIN_SECTIONS),
            "required_flow_stages": list(REQUIRED_FLOW_STAGES),
            "rendering": {
                "lesson_flow_as_table": True,
                "stage_columns": ["Етап", "Учитель", "Учні", "Завдання"],
            },
        }

    def _build_sources_context(self, source_files: list[str] | None, context: str) -> str:
        parts = [self._clean(context)]
        files = [item for item in (source_files or []) if item]
        if files and hasattr(self.legacy, "_build_sources_bundle"):
            try:
                bundle = self.legacy._build_sources_bundle(files, blueprint=None)
                bundle_context = self._clean(bundle.get("context"))
                if bundle_context:
                    parts.append(bundle_context)
            except Exception:
                pass
        return "\n\n".join(part for part in parts if part)

    @staticmethod
    def _escape_control_chars_in_json_strings(value: str) -> str:
        result: list[str] = []
        in_string = False
        escaped = False
        for char in value:
            if escaped:
                result.append(char)
                escaped = False
                continue
            if char == "\\":
                result.append(char)
                escaped = True
                continue
            if char == '"':
                result.append(char)
                in_string = not in_string
                continue
            if in_string and char == "\n":
                result.append("\\n")
                continue
            if in_string and char == "\r":
                result.append("\\r")
                continue
            if in_string and char == "\t":
                result.append("\\t")
                continue
            if in_string and ord(char) < 32:
                continue
            result.append(char)
        return "".join(result)

    @staticmethod
    def _json_from_model_text(text: str) -> dict:
        value = str(text or "").strip()
        value = re.sub(r"^\s*```(?:json)?\s*", "", value, flags=re.IGNORECASE)
        value = re.sub(r"\s*```\s*$", "", value)
        start = value.find("{")
        end = value.rfind("}")
        if start >= 0 and end > start:
            value = value[start : end + 1]
        try:
            data = json.loads(value)
        except json.JSONDecodeError:
            data = json.loads(LessonGeneratorV25._escape_control_chars_in_json_strings(value))
        if not isinstance(data, dict):
            raise ValueError("Model returned JSON that is not an object")
        return data

    @staticmethod
    def _as_list(value: Any) -> list[str]:
        if isinstance(value, list):
            return [re.sub(r"\s+", " ", str(item or "")).strip() for item in value if str(item or "").strip()]
        if isinstance(value, str) and value.strip():
            return [line.strip() for line in value.splitlines() if line.strip()]
        return []

    @classmethod
    def _normalize_lesson(cls, data: dict, *, topic: str, grade: str, subject: str) -> dict:
        goals = data.get("goals") if isinstance(data.get("goals"), dict) else {}
        flow = data.get("lesson_flow") if isinstance(data.get("lesson_flow"), list) else []
        normalized_flow = []
        flow_by_title = {
            cls._clean(item.get("stage") or item.get("title")).lower(): item
            for item in flow
            if isinstance(item, dict)
        }
        flow_by_key = {
            cls._stage_key(cls._clean(item.get("stage") or item.get("title"))): item
            for item in flow
            if isinstance(item, dict)
        }
        for stage_title in REQUIRED_FLOW_STAGES:
            raw = flow_by_title.get(stage_title.lower()) or flow_by_key.get(cls._stage_key(stage_title)) or {}
            content = cls._as_list(raw.get("content") or raw.get("steps") or raw.get("activities"))
            if not content:
                content = [
                    item
                    for item in (
                        cls._clean(raw.get("teacher")),
                        cls._clean(raw.get("students")),
                        cls._clean(raw.get("task") or raw.get("exercise")),
                    )
                    if item
                ]
            normalized_flow.append(
                {
                    "stage": stage_title,
                    "teacher": cls._clean(raw.get("teacher")),
                    "students": cls._clean(raw.get("students")),
                    "task": cls._clean(raw.get("task") or raw.get("exercise")),
                    "content": content,
                }
            )
        return {
            "topic": cls._clean(data.get("topic") or topic),
            "grade": cls._clean(data.get("grade") or grade),
            "subject": cls._clean(data.get("subject") or subject),
            "lesson_type": cls._clean(data.get("lesson_type") or "Урок вивчення і первинного закріплення нового матеріалу"),
            "goals": {
                "навчальна": cls._clean(goals.get("навчальна") or goals.get("educational")),
                "розвивальна": cls._clean(goals.get("розвивальна") or goals.get("developmental")),
                "виховна": cls._clean(goals.get("виховна") or goals.get("educational_value") or goals.get("value")),
            },
            "expected_results": cls._as_list(data.get("expected_results")),
            "equipment": cls._as_list(data.get("equipment") or data.get("materials")),
            "lesson_flow": normalized_flow,
            "homework": cls._as_list(data.get("homework")),
        }

    @staticmethod
    def _lesson_to_text(lesson: dict) -> str:
        lines = [
            "Тема уроку",
            str(lesson.get("topic") or ""),
            "Клас",
            str(lesson.get("grade") or ""),
            "Предмет",
            str(lesson.get("subject") or ""),
            "Тип уроку",
            str(lesson.get("lesson_type") or ""),
            "Мета уроку",
        ]
        for label, value in (lesson.get("goals") or {}).items():
            if value:
                lines.append(f"{label}: {value}")
        lines.append("Очікувані результати")
        lines.extend(lesson.get("expected_results") or [])
        lines.append("Обладнання та матеріали")
        lines.extend(lesson.get("equipment") or [])
        lines.append("Хід уроку")
        for stage in lesson.get("lesson_flow") or []:
            lines.append(stage.get("stage") or "")
            content = stage.get("content") or []
            if content:
                lines.extend(content)
            else:
                lines.append(stage.get("teacher") or "")
                lines.append(stage.get("students") or "")
                lines.append(stage.get("task") or "")
        lines.append("Домашнє завдання")
        lines.extend(lesson.get("homework") or [])
        return normalize_lesson_text("\n".join(line for line in lines if str(line).strip()))

    @staticmethod
    def _qa_lesson(lesson: dict, *, topic: str, subject: str) -> dict:
        text = LessonGeneratorV25._lesson_to_text(lesson)
        validation = assess_lesson_text(text, topic=topic, subject=subject)
        issues: list[str] = []
        if not lesson.get("topic"):
            issues.append("missing_topic")
        if not all((lesson.get("goals") or {}).get(key) for key in ("навчальна", "розвивальна", "виховна")):
            issues.append("missing_goal_parts")
        if len(lesson.get("expected_results") or []) < 3:
            issues.append("expected_results_too_short")
        if len(lesson.get("equipment") or []) < 2:
            issues.append("equipment_too_short")
        flow = lesson.get("lesson_flow") or []
        if len(flow) != len(REQUIRED_FLOW_STAGES):
            issues.append("lesson_flow_stage_count_invalid")
        for stage in flow:
            if len(stage.get("content") or []) < 2 and (
                not stage.get("teacher") or not stage.get("students") or not stage.get("task")
            ):
                issues.append(f"stage_incomplete:{stage.get('stage')}")
        if len(lesson.get("homework") or []) < 1:
            issues.append("homework_missing")
        issues.extend(validation.get("reasons") or [])
        return {
            "text": text,
            "validation": validation,
            "issues": sorted(set(issues)),
            "ok": not issues,
        }

    @staticmethod
    def _subject_rules(subject: str) -> str:
        low = str(subject or "").lower()
        if "математ" in low:
            return "Математика: кожен змістовний етап містить конкретні числа, вирази, рівності або короткі задачі за темою."
        if "мова" in low:
            return "Українська мова: кожен змістовний етап містить конкретні слова, словосполучення, речення, мінітекст або мовне спостереження за темою."
        if "літера" in low or "читан" in low:
            return "Читання/література: кожен змістовний етап містить роботу з текстом, героєм, читацьким питанням або творчим завданням."
        if "ядс" in low or "дослідж" in low:
            return "ЯДС: кожен змістовний етап містить спостереження, дослід, практичну дію, роботу з реальною ситуацією або висновок."
        return "Кожен змістовний етап має містити конкретний навчальний матеріал за темою."

    def _build_generation_prompt(
        self,
        *,
        topic: str,
        grade: str,
        subject: str,
        requirements: str,
        source_context: str,
        blueprint: dict,
    ) -> tuple[str, str]:
        system_instruction = (
            "Ти досвідчений український методист НУШ для початкової школи. "
            "Створи повний, цілісний конспект уроку за структурою шаблону. "
            "Не збирай урок з випадкових фрагментів: продумай педагогічну логіку від мотивації до підсумку."
        )
        prompt = f"""
Згенеруй повний конспект уроку НУШ українською мовою.

Вхідні дані:
- Предмет: {subject}
- Клас: {grade}
- Тема: {topic}
- Побажання вчителя: {requirements or "немає"}
- Додатковий контекст/матеріали: {source_context or "немає"}

Blueprint шаблону, витягнутий з DOCX:
{json.dumps(blueprint, ensure_ascii=False, indent=2)}

Предметне правило:
{self._subject_rules(subject)}

Вимоги:
1. Генеруй ОДИН цілісний урок, а не окремі незалежні шматки.
2. Дотримуйся структури: Тема уроку, Клас, Предмет, Тип уроку, Мета уроку, Очікувані результати, Обладнання та матеріали, Хід уроку, Домашнє завдання.
3. У "Хід уроку" мають бути всі етапи з blueprint.required_flow_stages у цьому порядку.
4. Кожен етап має містити content: 3-6 послідовних абзаців ходу уроку. Не діли конспект на колонки чи ролі "учитель/учні".
5. Не копіюй референсний шаблон дослівно; використовуй його як модель структури й стилю.
6. Не пиши загальні фрази без змісту: "виконати завдання", "провести бесіду", "обговорити тему", "формувати вміння", "ознайомити учнів", "закріпити знання", "розвивати мислення", "виховувати інтерес".
7. Якщо тема вузька, усі ключові вправи в актуалізації та закріпленні мають тренувати саме заявлену тему.
8. Мета уроку має називати конкретну дію за темою, наприклад "складати таблицю множення на 6", а не загальне "формувати вміння".
9. Домашнє завдання має бути конкретним: подай 2-4 точні приклади, задачу або сторінку/номер без фрази "виконати завдання".
10. У значеннях JSON не починай рядки з "-", "•", "*", "1.", "2." або інших маркерів списку. Пиши звичайні речення без markdown.

Поверни лише валідний JSON такого формату:
{{
  "topic": "...",
  "grade": "...",
  "subject": "...",
  "lesson_type": "...",
  "goals": {{
    "навчальна": "...",
    "розвивальна": "...",
    "виховна": "..."
  }},
  "expected_results": ["...", "...", "..."],
  "equipment": ["...", "..."],
  "lesson_flow": [
    {{"stage": "Організаційний момент", "content": ["Звичайне речення без маркера списку.", "Ще одне речення без дефіса на початку."]}},
    {{"stage": "Актуалізація опорних знань", "content": ["...", "...", "..."]}},
    {{"stage": "Мотивація навчальної діяльності", "content": ["...", "...", "..."]}},
    {{"stage": "Вивчення нового матеріалу", "content": ["...", "...", "..."]}},
    {{"stage": "Закріплення знань", "content": ["...", "...", "..."]}},
    {{"stage": "Підсумок уроку", "content": ["...", "...", "..."]}}
  ],
  "homework": ["..."]
}}
""".strip()
        return system_instruction, prompt

    def _build_repair_prompt(
        self,
        *,
        lesson: dict,
        topic: str,
        grade: str,
        subject: str,
        requirements: str,
        issues: list[str],
        blueprint: dict,
    ) -> tuple[str, str]:
        system_instruction = (
            "Ти редактор НУШ-конспектів. Виправ тільки методичні й структурні проблеми, "
            "зберігаючи цілісність уроку та JSON-формат."
        )
        prompt = f"""
Виправ конспект за списком проблем.

Предмет: {subject}
Клас: {grade}
Тема: {topic}
Побажання: {requirements or "немає"}
Проблеми QA: {json.dumps(issues, ensure_ascii=False)}
Blueprint: {json.dumps(blueprint, ensure_ascii=False)}

Поточний JSON:
{json.dumps(lesson, ensure_ascii=False, indent=2)}

Правила виправлення:
- Прибери загальні фрази: "виконати завдання", "формувати вміння", "ознайомити учнів", "закріпити знання", "розвивати мислення", "виховувати інтерес".
- У меті, очікуваних результатах, актуалізації, вивченні нового матеріалу, закріпленні та домашньому завданні прямо назви тему "{topic}" або її ключову дію.
- У кожному етапі використовуй content: 3-6 послідовних абзаців живого ходу уроку без колонок і без розділення на ролі teacher/students/task.
- Домашнє завдання запиши конкретно: приклади, коротка задача, тренування таблиці/правила за темою.
- У рядках content, expected_results, equipment і homework не використовуй маркери списку на початку: "-", "•", "*", "1.", "2.". Пиши чистий текст.

Поверни лише повний валідний JSON у тому самому форматі. Не додавай пояснень.
""".strip()
        return system_instruction, prompt

    async def _call_json_model(
        self,
        *,
        call_name: str,
        system_instruction: str,
        prompt: str,
        runtime_trace: dict,
        timeout_sec: float,
        temperature: float,
    ) -> dict:
        response = await self.legacy._execute_model_call(
            call_name=call_name,
            prompt=prompt,
            system_instruction=system_instruction,
            temperature=temperature,
            runtime_trace=runtime_trace,
            timeout_sec=timeout_sec,
            response_mime_type="application/json",
        )
        text = str(getattr(response, "text", "") if response else "").strip()
        if not text:
            error = ""
            for item in reversed(runtime_trace.get("model_calls") or []):
                if isinstance(item, dict) and item.get("name") == call_name:
                    error = str(item.get("error") or "")
                    break
            raise RuntimeError(f"{call_name} did not return JSON lesson text" + (f": {error}" if error else ""))
        return self._json_from_model_text(text)

    @staticmethod
    def _set_cell_text(cell, text: str, *, bold: bool = False):
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(str(text or ""))
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.bold = bold

    @staticmethod
    def _set_table_borders(table):
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                borders = tc_pr.first_child_found_in("w:tcBorders")
                if borders is None:
                    borders = OxmlElement("w:tcBorders")
                    tc_pr.append(borders)
                for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                    element = borders.find(qn(f"w:{edge}"))
                    if element is None:
                        element = OxmlElement(f"w:{edge}")
                        borders.append(element)
                    element.set(qn("w:val"), "single")
                    element.set(qn("w:sz"), "4")
                    element.set(qn("w:space"), "0")
                    element.set(qn("w:color"), "808080")

    @staticmethod
    def _add_paragraph(
        doc: Document,
        text: str,
        *,
        bold: bool = False,
        italic: bool = False,
        size: int = 14,
        align=None,
        space_after: int = 4,
        left_indent: int = 0,
        first_line_indent: int = 0,
    ):
        paragraph = doc.add_paragraph()
        if align is not None:
            paragraph.paragraph_format.alignment = align
        paragraph.paragraph_format.space_after = Pt(space_after)
        paragraph.paragraph_format.left_indent = Pt(left_indent)
        paragraph.paragraph_format.first_line_indent = Pt(first_line_indent)
        run = paragraph.add_run(str(text or ""))
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        return paragraph

    @staticmethod
    def _add_heading(doc: Document, text: str, *, level: int = 1):
        size = 16 if level == 1 else 14
        paragraph = LessonGeneratorV25._add_paragraph(
            doc,
            text,
            bold=True,
            size=size,
            space_after=6 if level == 1 else 4,
        )
        return paragraph

    @staticmethod
    def _add_plain_item(doc: Document, text: str):
        value = str(text or "").strip()
        if not value:
            return None
        return LessonGeneratorV25._add_paragraph(doc, value, size=13, left_indent=14, first_line_indent=0)

    @staticmethod
    def _stage_number(index: int) -> str:
        numerals = ("I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X")
        return numerals[index] if 0 <= index < len(numerals) else str(index + 1)

    @staticmethod
    def _add_labeled_block(doc: Document, label: str, text: str):
        value = str(text or "").strip()
        if not value:
            return
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.left_indent = Pt(14)
        label_run = paragraph.add_run(f"{label}: ")
        label_run.font.name = "Times New Roman"
        label_run.font.size = Pt(13)
        label_run.bold = True
        text_run = paragraph.add_run(value)
        text_run.font.name = "Times New Roman"
        text_run.font.size = Pt(13)

    def _render_docx(self, lesson: dict, *, output_path: Path):
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Pt(56)
        section.bottom_margin = Pt(56)
        section.left_margin = Pt(56)
        section.right_margin = Pt(56)

        self._add_paragraph(
            doc,
            "Конспект уроку",
            bold=True,
            size=16,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=10,
        )
        self._add_labeled_block(doc, "Тема", lesson.get("topic") or "")
        self._add_labeled_block(doc, "Клас", lesson.get("grade") or "")
        self._add_labeled_block(doc, "Предмет", lesson.get("subject") or "")
        self._add_labeled_block(doc, "Тип уроку", lesson.get("lesson_type") or "")

        self._add_heading(doc, "Мета уроку")
        for label, value in (lesson.get("goals") or {}).items():
            if value:
                self._add_labeled_block(doc, label.capitalize(), value)

        self._add_heading(doc, "Очікувані результати")
        for item in lesson.get("expected_results") or []:
            self._add_plain_item(doc, item)

        self._add_heading(doc, "Обладнання та матеріали")
        for item in lesson.get("equipment") or []:
            self._add_plain_item(doc, item)

        self._add_heading(doc, "Хід уроку")
        for idx, stage in enumerate(lesson.get("lesson_flow") or []):
            title = stage.get("stage") or ""
            self._add_paragraph(
                doc,
                f"{self._stage_number(idx)}. {title}",
                bold=True,
                size=14,
                space_after=3,
            )
            content = stage.get("content") or []
            if not content:
                content = [
                    item
                    for item in (stage.get("teacher"), stage.get("students"), stage.get("task"))
                    if str(item or "").strip()
                ]
            for item in content:
                self._add_plain_item(doc, item)

        self._add_heading(doc, "Домашнє завдання")
        for item in lesson.get("homework") or []:
            self._add_plain_item(doc, item)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

    async def generate_lesson_files(
        self,
        *,
        topic: str,
        grade: str,
        requirements: str,
        mode: str,
        subject: str,
        context: str = "",
        source_file: str | None = None,
        source_files: list[str] | None = None,
        request_id: str | None = None,
    ) -> dict:
        if mode != "docx":
            raise ValueError("v2.5 currently supports DOCX generation only")

        request_id = request_id or secrets.token_hex(6)
        started = time.perf_counter()
        diagnostics: dict[str, Any] = {
            "request_id": request_id,
            "pipeline_version": self.pipeline_version,
            "status": "started",
            "model_calls": [],
            "model_calls_total": 0,
            "refinement_used": False,
            "references": [],
            "references_count": 0,
            "validation": {},
        }
        runtime_trace = {"request_id": request_id, "model_calls": diagnostics["model_calls"]}

        template_path, parsed_template = self._select_template(subject=subject, grade=grade, topic=topic)
        blueprint = self._build_blueprint(parsed_template, template_path, subject=subject)
        diagnostics["references"] = [template_path.name] if template_path else []
        diagnostics["references_count"] = len(diagnostics["references"])
        diagnostics["template_blueprint"] = blueprint

        normalized_source_files = [item for item in (source_files or ([source_file] if source_file else [])) if item]
        source_context = self._build_sources_context(normalized_source_files, context)

        sys_instr, prompt = self._build_generation_prompt(
            topic=topic,
            grade=grade,
            subject=subject,
            requirements=requirements,
            source_context=source_context,
            blueprint=blueprint,
        )
        timeout_1 = float(os.getenv("GENERATOR_V25_TIMEOUT_SEC", os.getenv("GENERATOR_REWRITE_TIMEOUT_SEC", "55")))
        raw_lesson = await self._call_json_model(
            call_name="v25_generate_full_lesson",
            system_instruction=sys_instr,
            prompt=prompt,
            runtime_trace=runtime_trace,
            timeout_sec=timeout_1,
            temperature=0.35,
        )
        lesson = self._normalize_lesson(raw_lesson, topic=topic, grade=grade, subject=subject)
        qa = self._qa_lesson(lesson, topic=topic, subject=subject)
        diagnostics["validation"] = qa["validation"]
        diagnostics["qa_issues"] = qa["issues"]
        diagnostics["qa_ok"] = qa["ok"]

        repair_enabled = os.getenv("GENERATOR_V25_REPAIR", "1").strip().lower() in {"1", "true", "yes", "on"}
        if repair_enabled and qa["issues"]:
            r_sys, r_prompt = self._build_repair_prompt(
                lesson=lesson,
                topic=topic,
                grade=grade,
                subject=subject,
                requirements=requirements,
                issues=qa["issues"],
                blueprint=blueprint,
            )
            timeout_2 = float(os.getenv("GENERATOR_V25_REPAIR_TIMEOUT_SEC", os.getenv("GENERATOR_REWRITE_REFINE_TIMEOUT_SEC", "35")))
            repaired_raw = await self._call_json_model(
                call_name="v25_repair_full_lesson",
                system_instruction=r_sys,
                prompt=r_prompt,
                runtime_trace=runtime_trace,
                timeout_sec=timeout_2,
                temperature=0.2,
            )
            repaired = self._normalize_lesson(repaired_raw, topic=topic, grade=grade, subject=subject)
            repaired_qa = self._qa_lesson(repaired, topic=topic, subject=subject)
            if len(repaired_qa["issues"]) <= len(qa["issues"]):
                lesson = repaired
                qa = repaired_qa
                diagnostics["refinement_used"] = True
                diagnostics["validation_after_refinement"] = repaired_qa["validation"]
                diagnostics["qa_issues_after_refinement"] = repaired_qa["issues"]
                diagnostics["qa_ok"] = repaired_qa["ok"]

        file_id = f"{secrets.token_hex(4)}_{sanitize_filename(str(grade) or 'grade', default='grade')}"
        output_path = Path("storage") / f"Lesson_v25_{file_id}.docx"
        self._render_docx(lesson, output_path=output_path)

        diagnostics["lesson_json"] = lesson
        diagnostics["lesson_text"] = qa["text"]
        diagnostics["model_calls_total"] = len(diagnostics["model_calls"])
        diagnostics["status"] = "success"
        diagnostics["duration_ms"] = int((time.perf_counter() - started) * 1000)
        diagnostics["output_name"] = output_path.name
        diagnostics["output_path"] = str(output_path)
        diagnostics["output_ext"] = ".docx"
        diagnostics["output_size_bytes"] = output_path.stat().st_size if output_path.exists() else 0
        diagnostics["output_oversize"] = False

        return {
            "path": str(output_path),
            "name": output_path.name,
            "diagnostics": diagnostics,
        }
